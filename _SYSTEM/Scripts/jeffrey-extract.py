#!/usr/bin/env python3
"""jeffrey-extract.py — extract plain text from a PDF / Word / Excel file for Jeffrey's local index.

Prints extracted text to stdout (capped). On ANY error prints nothing, so the indexer falls back to a
filename-only row. Deps (Jeffrey venv): pdfminer.six, python-docx, openpyxl.
Usage: python jeffrey-extract.py <file.pdf|file.docx|file.xlsx>
"""
import sys, os

# Windows consoles default to cp1252, which can't encode text like '→' extracted from docs — force
# UTF-8 so the write never crashes (the node indexer reads stdout as utf8).
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

CAP = 200_000  # hard char cap — plenty for FTS matching + snippets, bounds runaway extraction


def extract_pdf(p):
    from pdfminer.high_level import extract_text
    return extract_text(p) or ""


def extract_docx(p):
    import docx
    d = docx.Document(p)
    parts = [para.text for para in d.paragraphs if para.text]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def extract_xlsx(p):
    import openpyxl
    wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
    out, total = [], 0
    try:
        for ws in wb.worksheets:
            out.append(f"# {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    line = "\t".join(cells)
                    out.append(line); total += len(line)
                    if total > CAP:
                        return "\n".join(out)
    finally:
        wb.close()
    return "\n".join(out)


def main():
    if len(sys.argv) < 2:
        return
    p = sys.argv[1]
    ext = os.path.splitext(p)[1].lower()
    try:
        if ext == ".pdf":
            txt = extract_pdf(p)
        elif ext == ".docx":
            txt = extract_docx(p)
        elif ext == ".xlsx":
            txt = extract_xlsx(p)
        else:
            return
    except Exception:
        return  # silent → indexer keeps the filename-only row
    sys.stdout.write((txt or "")[:CAP])


if __name__ == "__main__":
    main()
